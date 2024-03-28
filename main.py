from docx import Document
import targetList


docxFilePath = r"C:\Users\wilso\Desktop\syllabus\Syllabus - Winter 2024 - Relational Databases.docx"

doc = Document(docxFilePath)



relatedText=[]
'''
Step1: return[['Week 1',' '], ['Week 2',' ']]
'''
def weekFinder():
    all_related_text = []  # 用于存储所有行的相关文本列表
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if 'week' in cell.text.lower():
                    related_text = []  # 用于存储当前行的相关文本
                    for cell_in_row in row.cells:
                        related_text.append(cell_in_row.text.strip())  # 将当前行的文本添加到列表中
                    all_related_text.append(related_text)  # 将当前行的相关文本列表添加到总列表中
    return all_related_text
    #回傳多個列表

def dupCheck(list):
    noDupList=[]
    for sublist in list:
        if sublist not in noDupList:
            noDupList.append(sublist)
        if noDupList[0][0].lower() != 'week 1':
            noDupList.remove(noDupList[0])
    return noDupList
'''
Step2: 
'''
finalList = []


noDupList=dupCheck(weekFinder())
print(noDupList)




'''
doc = Document()
table = doc.add_table(rows=15, columns=1)

for i in range(len(finalList)):
    table.cell(i,0).text = finalList[i]
    def splitList(list, keywords):
    subLists = []
    subList = []
    for item in list:
        if any(keyword in item for keyword in keywords):
            if subList:
                subLists.append(subList)
                subList = []
        subList.append(item)
    if subList:
        subLists.append(subList)
    return subLists
    
    
    def targetChecker(list):
    for item in list:
        for word in item.split():
            if word in targetList.weekList:
                finalList.append(word)
                if word in targetList.targetList:
                    finalList.append(word)
    return finalList
    
    
    def print_paragraphs(doc):
    for para in doc.paragraphs:
        print(para.text)
    '''
