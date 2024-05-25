from docx import Document
from docx.enum.section import WD_ORIENTATION
from constants import targetList
import os, re

'''
Step1: 
Using two functions to return a nested list with a length of 15.
The first function iterates through all the tables in the document file. 
Once it finds "week" in a cell, it stores the contents of the same row in list format. 
Finally, it stores everything into a nested list and returns it.
The second function removes duplicate lists and ensures that the first sublist begins with "week 1". 
It returns a nested list sorted by week, with a length of 15.
return[['Week 1',' '], ['Week 2',' ']]
'''

def weekFinder(docxFilePath):
    # Open the specified Word document file.
    doc = Document(docxFilePath)
    # Initialize an empty list to store all related text.
    allRelatedText = []
    columnAmounts = []
    # Initialize table variable to be tracked
    tables = 0
    # Store a regex pattern to search specifically for the word "week"
    # We use regex so that it wont pick up "weekly" (or anything else with week in it) but will still pick up the word "weeK" even if theres no spaces
    pattern = r'\bweek\b'
    # Iterate through each table in the document.
    for table in doc.tables:
        # Store the first cell in the table
        firstCell = table.cell(0, 0).text.strip().lower()
        # Store columns amount for later comparison
        columnAmounts.append(len(table.columns))
        # This will run if the first cell contains "week" or "module" or if tables = 1 (meaning that we've already gone through one table)
        # We check the tables as sometimes the weekly table will be split across pages causing it to be multiple tables
        # We check the column amounts to make sure its the same table extended as some syllabus' have special cases that require extra checks
        # So by checking if tables = 1 and the column amounts are the same we ensure that we've grabbed the data from one page and now grab the rest from the other
        if (re.search(pattern, firstCell) or "module" in firstCell) or (
                tables == 1 and all(element == columnAmounts[0] for element in columnAmounts)):
            # Increment tables
            tables += 1
            # Iterate through each row in tables
            for row in table.rows:
                # Initialize an empty list to store related text in the current row.
                relatedText = []

                # Iterate through each cell in the row to collect related text.
                for cellInRow in row.cells:
                    # Append the stripped text of each cell to the relatedText list.
                    relatedText.append(cellInRow.text.strip())
                # Append the relatedText list of the current row to the allRelatedText list.
                allRelatedText.append(relatedText)
                print(allRelatedText)
        # If the table is not the timetable, remove its row amount
        else:
            columnAmounts.pop(-1)

    # Return the list containing related text from all rows.
    return allRelatedText


def dupCheck(listOne):
    # Initialize an empty list to store unique sublists.
    noDupList = []
    # Iterate through each sublist in the input list.
    for sublist in listOne:
        # Check if the sublist is not already in the noDupList.
        if sublist not in noDupList:
            # Check if the list is not made up of all empty strings
            if not all(string == "" for string in sublist):
                # If not, append it to the noDupList.
                noDupList.append(sublist)

    # If list is not empty (information found)
    if len(noDupList) > 0:
        # Delete the first sublist as it is column headers and useless to us
        noDupList.remove(noDupList[0])
    # Return the list with duplicates removed and only the first sublist being 'week 1'.
    return noDupList


'''
Step2: 
Using two functions to organize a nested list with a length of 15 into a single list containing only keywords.
The first function iterates through the entire nested list, replicating the method from step 1. 
It extracts keywords and adds them to a sublist. Then, it adds these 15 sublists to a larger nested list. 
Finally, it returns a nested list of length 15 containing only keywords.
The second function uses .join to merge the sublists of the nested list into strings. 
Ultimately, it returns a single list with a length of 15.
'''


def targetWordsChecker(listOne):
    # Initialize an empty list to store the modified sublists.
    newList = []
    # Create a regular expression pattern that matches any word in the targetList.
    pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, targetList.targetList)) + r')', re.IGNORECASE)
    # Iterate through each sublist in the input list.
    for i in range(len(listOne)):
        # Initialize an empty list to store modified words in the current sublist.
        modifiedSublist = []
        # Iterate through each word in the current sublist.
        for j in range(len(listOne[i])):
            # Split the current word into a list of words.
            #words = listOne[i][j].split()
            # Filter out words that are not in the targetList.
            #modifiedWords = [word for word in words if word.strip().lower() in targetList.targetList]
            # Join the modified words back into a string and append to the modifiedSublist.
            #modifiedSublist.append(' '.join(modifiedWords))
            # Find all occurrences of the pattern in the current word.
            matches = pattern.findall(listOne[i][j])
            # Append the matches to the modifiedSublist.
            modifiedSublist.extend(matches)
            print(modifiedSublist)
        # Append the modifiedSublist to the newList.
        newList.append(modifiedSublist)
    # Return the newList containing modified sublists.
    return newList


def concatList(listOne):
    # Initialize an empty list to store concatenated strings.
    result = []
    # Iterate through each sublist in the input list.
    for i in range(len(listOne)):
        # Join the elements of the current sublist into a single string and append to result.
        result.append(' '.join(listOne[i]))
    # Return the list containing concatenated strings.
    return result


'''
Step3: 
Inserting a list of length fifteen into the second column of a table, starting from the second row, 
in a docx file, and saving it as a new file.
'''

# Change the file from portrait to landscape if needed
def setLandscape(doc):
    section = doc.sections[0]
    if section.orientation != WD_ORIENTATION.LANDSCAPE:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        newWidth, newHeight = section.page_height, section.page_width
        section.page_width = newWidth
        section.page_height = newHeight

def insertToTable(listOne, outputFilePath, column, filePath):
    # For the first file use the template
    if column == 1:
        # Open the Word document file provided by client.
        doc = Document(r"static/schedule.docx")
    # For every other file use the created word file
    else:
        # Open the Word document created
        doc = Document(outputFilePath)
    # Make the file landscape orientation
    setLandscape(doc)
    # Get the first table in the document.
    table = doc.tables[0]
    # Include program name (syllabus file name) in the first row
    table.cell(0, column).text = os.path.basename(filePath)
    # Iterate through each element in the input list.
    for i in range(len(listOne)):
        # Get the cell in the column based off the file and i+1 row of the table.
        # If the table pulled from has more than 15 rows, dont run the rest as it will overflow the schedule table
        # (No syllabus should have more than 15 unless it's by mistake)
        if i + 1 <= 15:
            cell = table.cell(i + 1, column)

        # Set the text of the cell to the corresponding element in the input list.
        cell.text = listOne[i]

    # Save the modified document to a new file.
    doc.save(outputFilePath)