from docx import Document
from constants import targetList

'''
Step1: 
Using 3 functions to return a nested list with a length of 15.

findWeekInRow(row): This function iterates through all cells in a table row. Once it finds "week" in a cell, it stores 
the contents of the same row in list format and returns it. 

tableInfoExtract(docxFilePath): This function iterates through all tables in the document. For each row in each table, 
it calls the findWeekInRow(row) function to get related text. Finally, it stores all related text into a nested list 
and returns it.

dupCheck(listOne): This function removes duplicate lists and ensures that the first sublist begins with "week 1". 
It returns a nested list sorted by week, with a length of 15.

return[['Week 1',' '], ['Week 2',' ']....['Week 15',' ']]
'''
def weekFinder(row):
    # Initialize an empty list to store related text in the current row.
    allRelatedText = []
    # Iterate through each cell in the row.
    for cell in row.cells:
        # Check if the text in the cell contains 'week' (case insensitive).
        if 'week' in cell.text.lower():
            # Initialize an empty list to store related text in the current row.
            relatedText = []

            # Iterate through each cell in the row again to collect related text.
            for cellInRow in row.cells:
                # Append the stripped text of each cell to the relatedText list.
                relatedText.append(cellInRow.text.strip())

            # Append the relatedText list of the current row to the allRelatedText list.
            allRelatedText.append(relatedText)

    # Return the list containing related text from the current row.
    return allRelatedText

def tableInfoExtract(docxFilePath):
    # Open the specified Word document file.
    doc = Document(docxFilePath)

    # Initialize an empty list to store all related text.
    allRelatedText = []

    # Use list comprehension to collect all related text.
    allRelatedText += [listOne for table in doc.tables for row in table.rows for listOne in weekFinder(row)]

    # Return the list containing related text from all rows.
    return allRelatedText


def dupCheck(listOne):
    # Initialize an empty list to store unique sublists.
    noDupList = []

    # Use list comprehension to remove duplicates from the input list.
    noDupList += [sublist for sublist in listOne if sublist not in noDupList]

    # Check if the first sublist's first element is not 'week 1' (case insensitive).
    if noDupList and noDupList[0][0].lower() != 'week 1':
        # If it's not 'week 1', remove the first sublist from the noDupList.
        noDupList.remove(noDupList[0])

    # Return the list with duplicates removed and only the first sublist being 'week 1'.
    return noDupList


'''
Step2: 
Using two functions to organize a nested list with a length of 15 into a single list containing only keywords.

The first function iterates through the entire nested list, replicating the method from step 1. 
It extracts keywords and adds them to a sublist. Then, it adds these 15 sublists to a larger nested list. 
Finally, it returns a nested list of length 15 containing only keywords.

The second function uses .join to merge the sublists of the nested list into strings. 
Ultimately, it returns a single list with a length of 15.
'''


def targetWordsChecker(listOne):
    # Initialize an empty list to store the modified sublists.
    newList = []

    # Iterate through each sublist in the input list.
    for i in range(len(listOne)):
        # Initialize an empty list to store modified words in the current sublist.
        modifiedSublist = []

        # Iterate through each word in the current sublist.
        for j in range(len(listOne[i])):
            # Split the current word into a list of words.
            words = listOne[i][j].split()

            # Filter out words that are not in the targetList.
            modifiedWords = [word for word in words if word.lower() in targetList.targetList]

            # Join the modified words back into a string and append to the modifiedSublist.
            modifiedSublist.append(' '.join(modifiedWords))

        # Append the modifiedSublist to the newList.
        newList.append(modifiedSublist)

    # Return the newList containing modified sublists.
    return newList


def concatList(listOne):
    # Initialize an empty list to store concatenated strings.
    result = []
    # Use list comprehension to concatenate the strings in each sublist.
    result += [''.join(listOne[i]) for i in range(len(listOne))]

    # Return the list containing concatenated strings.
    return result

'''
Step3: 
Inserting a list of length fifteen into the second column of a table, starting from the second row, 
in a docx file, and saving it as a new file.
'''

def insertIntoTable(listOne, outputFilePath):
    # Open the Word document file provided by client.
    doc = Document(r"static/schedule.docx")

    # Get the first table in the document.
    table = doc.tables[0]

    # Iterate through each element in the input list.
    for i in range(len(listOne)):
        # Get the cell in the first column and i+1 row of the table.
        cell = table.cell(i + 1, 1)

        # Set the text of the cell to the corresponding element in the input list.
        cell.text = listOne[i]

    # Save the modified document to a new file.
    doc.save(outputFilePath)

